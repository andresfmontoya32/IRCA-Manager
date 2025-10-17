# ============================================================
# Correspondencia masiva: Excel (TAGS) → Word (plantilla)
# ============================================================

from pathlib import Path
from datetime import datetime, date
import pandas as pd
from docx import Document
from tqdm.auto import tqdm
import re
from docx.shared import Inches

# ----------- DICCIONARIO DE FOTOS POR CIUDAD (PÉGALO AQUÍ) -----------
# Diccionario de fotos por ciudad - SEPTIEMBRE 2025
fotos_por_ciudad_septiembre = {
    "Aguachica": {
        "FOTO1": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Aguachica\202509_AP\Reg Foto\Muestras-P1.jpg",
        "FOTO2": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Aguachica\202509_AP\Reg Foto\Muestras-P2.jpg",
    },
    "Armenia": {
        "FOTO1": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Armenia\202509_AP\Reg Foto\Muestras-P1.jpg",
        "FOTO2": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Armenia\202509_AP\Reg Foto\Muestras-P2.jpg",
        "FOTO3": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Armenia\202509_AP\Reg Foto\Muestras-P3.jpg",
        "FOTO4": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Armenia\202509_AP\Reg Foto\Muestras-P4.jpg",
    },
    "Barranquilla": {
        "FOTO1": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\14.Calidad agua_ACTA 2_Consorcio\202509\202509_AP\Reg Foto\Muestras-P1.jpg",
        "FOTO2": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\14.Calidad agua_ACTA 2_Consorcio\202509\202509_AP\Reg Foto\Muestras-P2.jpg",
        "FOTO3": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\14.Calidad agua_ACTA 2_Consorcio\202509\202509_AP\Reg Foto\Muestras-P3.jpg",
        "FOTO4": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\14.Calidad agua_ACTA 2_Consorcio\202509\202509_AP\Reg Foto\Muestras-P4.jpg",
    },
    "Buenaventura": {
        "FOTO1": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Buenaventura\202509_AP\Reg Foto\Muestras-P1.jpg",
        "FOTO2": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Buenaventura\202509_AP\Reg Foto\Muestras-P2.jpg",
        "FOTO3": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Buenaventura\202509_AP\Reg Foto\Muestras-P3.jpg",
    },
    "Guapi": {
        "FOTO1": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Guapi\202509_AP\Reg Foto\Muestras-P1.jpg",
        "FOTO2": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Guapi\202509_AP\Reg Foto\Muestras-P2.jpg",
        "FOTO3": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Guapi\202509_AP\Reg Foto\Muestras-P3.jpg",
    },
    "Ipiales": {
        "FOTO1": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Ipiales\202509_AP\Reg Foto\Muestras-P1.jpg",
        "FOTO2": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Ipiales\202509_AP\Reg Foto\Muestras-P2.jpg",
        "FOTO3": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Ipiales\202509_AP\Reg Foto\Muestras-P3.jpg",
    },
    "Pasto": {
        "FOTO1": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Pasto\202509_AP\Reg Foto\Muestras-P1.jpg",
        "FOTO2": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Pasto\202509_AP\Reg Foto\Muestras-P2.jpg",
        "FOTO3": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Pasto\202509_AP\Reg Foto\Muestras-P3.jpg",
    },
    "Popayan": {
        "FOTO1": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Popayan\202509_AP\Reg Foto\Muestras-P1.jpg",
        "FOTO2": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Popayan\202509_AP\Reg Foto\Muestras-P2.jpg",
        "FOTO3": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Popayan\202509_AP\Reg Foto\Muestras-P3.jpg",
    },
    "Tolu": {
        "FOTO1": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Tolu\202509_AP\Reg Foto\Muestras-P1.jpg",
        "FOTO2": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Tolu\202509_AP\Reg Foto\Muestras-P2.jpg",
    },
    "Tumaco": {
        "FOTO1": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Tumaco\202509_AP\Reg Foto\Muestras-P1.jpg",
        "FOTO2": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Tumaco\202509_AP\Reg Foto\Muestras-P2.jpg",
        "FOTO3": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Tumaco\202509_AP\Reg Foto\Muestras-P3.jpg",
    },
    "San Andres": {
        "FOTO1": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\SAI\202509_AP\Reg Foto\Muestras-P1.jpg",
        "FOTO2": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\SAI\202509_AP\Reg Foto\Muestras-P2.jpg",
        "FOTO3": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\SAI\202509_AP\Reg Foto\Muestras-P3.jpg",
    },
    "Providencia": {
        "FOTO1": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Providencia\202509_AP\Reg Foto\Muestras-P1.jpg",
        "FOTO2": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Providencia\202509_AP\Reg Foto\Muestras-P2.jpg",
        "FOTO3": r"C:\Users\PAEROCIVIL\CONHYDRA SA ESP\Aerocivil - General\2. PIE\2. Hacer\12.Calidad Agua_ACTA 11\Providencia\202509_AP\Reg Foto\Muestras-P3.jpg",
    }
}

# ----------- 1. Ruta raíz que contiene las carpetas por aeropuerto -----------
ROOT_DIR = Path(r"C:\Users\PAEROCIVIL\Desktop\Automatizacion\AUTOMATIZACION\3.correspondencia\Datos")

# ----------- 2. Función para reemplazar tags en un documento ------------------
def reemplazar_tags(documento: Document, tags: dict):
    """
    Reemplaza {tag}, {{tag}} o <<tag>> en párrafos y celdas,
    conservando formato incluso si la etiqueta está partida en varios runs.
    También reemplaza en encabezados y pies de página.
    """
    variantes = []
    for t in tags:
        t_re = re.escape(t)
        variantes += [
            r"\{\{\s*" + t_re + r"\s*\}\}",   # {{tag}}
            r"\{\s*"  + t_re + r"\s*\}",      # {tag}
            r"<<\s*" + t_re + r"\s*>>"        # <<tag>>
        ]
    patron = re.compile("|".join(variantes), re.IGNORECASE)

    def _reemplazar_en_parrafo(parrafo):
        texto = "".join(run.text for run in parrafo.runs)
        if not patron.search(texto):
            return
        nuevo_texto = patron.sub(
            lambda m: tags[re.sub(r"[{}<> ]", "", m.group(0)).lower()], texto
        )
        # Vaciar runs y escribir el texto reemplazado en el primero
        for run in parrafo.runs:
            run.text = ""
        parrafo.runs[0].text = nuevo_texto

    # Párrafos fuera de tablas
    for p in documento.paragraphs:
        _reemplazar_en_parrafo(p)

    # Párrafos dentro de tablas
    for tabla in documento.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    _reemplazar_en_parrafo(p)

    # ----------- NUEVO: Reemplazo en encabezados y pies de página -----------
    for section in documento.sections:
        # Encabezado
        header = section.header
        for p in header.paragraphs:
            _reemplazar_en_parrafo(p)
        for tabla in header.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for p in celda.paragraphs:
                        _reemplazar_en_parrafo(p)
        # Pie de página (opcional, descomenta si lo necesitas)
        # footer = section.footer
        # for p in footer.paragraphs:
        #     _reemplazar_en_parrafo(p)
        # for tabla in footer.tables:
        #     for fila in tabla.rows:
        #         for celda in fila.cells:
        #             for p in celda.paragraphs:
        #                 _reemplazar_en_parrafo(p)

    return documento

# ----------- FUNCIÓN PARA INSERTAR FOTOS EN EL WORD ------------------
def insertar_fotos_en_docx(documento, ciudad, fotos_por_ciudad):
    """
    Inserta imágenes en el documento Word en los lugares donde hay tags de foto,
    usando el diccionario fotos_por_ciudad[ciudad].
    """
    if ciudad not in fotos_por_ciudad:
        print(f"⚠️  No hay fotos configuradas para la ciudad: {ciudad}")
        return documento

    fotos_ciudad = fotos_por_ciudad[ciudad]
    fotos_encontradas = 0
    fotos_no_encontradas = 0
    tags_procesados = 0
    
    print(f"🔍 Verificando fotos para {ciudad}:")
    
    # Verificar existencia de archivos antes de procesar
    for foto_tag, ruta_foto in fotos_ciudad.items():
        img_path = Path(ruta_foto)
        if img_path.exists():
            print(f"   ✅ {foto_tag}: Imagen encontrada")
        else:
            print(f"   ❌ {foto_tag}: Imagen NO encontrada - {ruta_foto}")
    
    # Buscar SOLO tags específicos de foto en las tablas del documento
    for tabla in documento.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    texto_original = p.text.strip()
                    
                    # Solo procesar si encuentra un tag ESPECÍFICO de foto
                    for i in range(1, 5):
                        foto_key = f"FOTO{i}"
                        # Diferentes formatos de tags posibles - SOLO TAGS ESPECÍFICOS
                        tag_patterns = [
                            "{FOTO" + str(i) + "}",
                            "{{FOTO" + str(i) + "}}",
                            "<<FOTO" + str(i) + ">>",
                            "{foto" + str(i) + "}",
                            "{{foto" + str(i) + "}}"
                        ]
                        
                        for tag_pattern in tag_patterns:
                            # SOLO si el texto ES EXACTAMENTE el tag (o solo contiene el tag)
                            if (tag_pattern == texto_original.lower() or 
                                tag_pattern in texto_original.lower()) and fotos_ciudad.get(foto_key):
                                
                                tags_procesados += 1
                                print(f"   🎯 Encontrado tag exacto: '{texto_original}' → procesando como {foto_key}")
                                
                                # Limpiar TODO el contenido del párrafo
                                p.clear()
                                
                                # Insertar la imagen si existe
                                img_path = Path(fotos_ciudad[foto_key])
                                if img_path.exists():
                                    # Tamaño prudente para caber en la celda de la tabla
                                    run = p.add_run()
                                    run.add_picture(str(img_path), width=Inches(2.0), height=Inches(1.5))
                                    fotos_encontradas += 1
                                    print(f"   📷 Imagen {foto_key} insertada correctamente (2.0x1.5 inches)")
                                else:
                                    p.add_run("Imagen no encontrada")
                                    fotos_no_encontradas += 1
                                    print(f"   ⚠️  Archivo no existe: {img_path}")
                                break  # Salir del loop de patterns una vez encontrado
    
    # SOLO si NO se encontraron tags específicos, informar
    if tags_procesados == 0:
        print("   ⚠️  No se encontraron tags específicos de fotos ({FOTO1}, {FOTO2}, etc.)")
        print("   💡 Asegúrate de que las celdas donde quieres fotos contengan tags como {FOTO1}, {FOTO2}, {FOTO3}, {FOTO4}")
    
    # Resumen de la inserción de fotos
    print(f"📊 Resumen fotos {ciudad}: {fotos_encontradas} insertadas | {fotos_no_encontradas} no encontradas | {tags_procesados} tags procesados")
    
    return documento

# ----------- 3. Formateo de valores ------------------------------------------
def fmt(val):
    """
    - Fechas/horas → 'YYYY-MM-DD' (solo fecha).
    - NaN → cadena vacía.
    - Resto → str(val)
    """
    if pd.isna(val):
        return ""
    if isinstance(val, (pd.Timestamp, datetime, date)):
        return val.strftime("%Y-%m-%d")            # cámbialo a "%d/%m/%Y" si prefieres
    return str(val)

# ----------- 4. Recorrer carpetas y procesar informes -------------------------
procesados, omitidos = 0, 0

print("\n🔎 Iniciando procesamiento de carpetas de aeropuertos...\n")

for carpeta in tqdm(list(ROOT_DIR.iterdir()), desc="Recorriendo aeropuertos"):
    if not carpeta.is_dir():
        continue

    ciudad = carpeta.name
    excel_path = carpeta / f"base_{ciudad}.xlsx"
    print(f"\n--- Procesando carpeta: {ciudad} ---")
    print(f"📁 Ruta carpeta: {carpeta}")
    print(f"📄 Buscando archivo Excel: {excel_path.name}")

    if not excel_path.exists():
        print(f"⚠️  No se encontró '{excel_path.name}'. Carpeta omitida.")
        omitidos += 1
        continue

    plantillas = list(carpeta.glob("*.docx"))
    print(f"📑 Buscando plantilla Word en carpeta...")
    if not plantillas:
        print(f"⚠️  Sin plantilla Word (.docx) en carpeta. Carpeta omitida.")
        omitidos += 1
        continue
    plantilla_path = plantillas[0]
    print(f"✅  Plantilla encontrada: {plantilla_path.name}")

    print(f"📊 Leyendo hoja 'TAGS' del Excel...")
    try:
        df_tags = pd.read_excel(excel_path, sheet_name="TAGS")
        print(f"✅  Hoja 'TAGS' leída correctamente.")
    except Exception as e:
        print(f"❌  Error leyendo hoja 'TAGS': {e}")
        omitidos += 1
        continue

    df_tags.columns = df_tags.columns.str.strip().str.upper()
    if {"ETIQUETA", "VALOR"} - set(df_tags.columns):
        print(f"❌  La hoja 'TAGS' no tiene columnas 'ETIQUETA' y 'VALOR'. Carpeta omitida.")
        omitidos += 1
        continue

    tags_dict = {
        k.strip().lower(): fmt(v)
        for k, v in zip(df_tags["ETIQUETA"], df_tags["VALOR"])
    }

    print(f"📝 Reemplazando tags en plantilla Word...")
    try:
        doc = Document(plantilla_path)
        doc = reemplazar_tags(doc, tags_dict)
        print(f"✅  Tags reemplazados correctamente.")

        print(f"🖼️ Insertando fotos en el documento Word...")
        doc = insertar_fotos_en_docx(doc, ciudad, fotos_por_ciudad_septiembre)
        salida_path = carpeta / f"reporte_{carpeta.name}.docx"
        doc.save(salida_path)
        print(f"✅  Documento generado: {salida_path.name}")
        procesados += 1
    except Exception as e:
        print(f"❌  Error al generar reporte: {e}")
        omitidos += 1

# ----------- 5. Resumen -------------------------------------------------------
print("\n==================== RESUMEN FINAL ====================")
print(f"🏁 Informes generados: {procesados}")
print(f"🚫 Carpetas omitidas: {omitidos}")
print("=======================================================\n")